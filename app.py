from flask import Flask, render_template, jsonify, request, send_file
from flask_caching import Cache
import psycopg2
import json
import csv
from datetime import datetime, timedelta
import os
import re
import subprocess
from pathlib import Path
from odf import text, teletype
from odf.opendocument import load, OpenDocumentText
from odf.table import Table, TableRow, TableCell
import base64
from PIL import Image
import io
from odt_editor import ODTEditor
from geocoder import geocode_location
import unicodedata
import zipfile
import tempfile
from docx import Document as DocxDocument
from docx.oxml.ns import qn

# Load city to state mapping for pheno_new historical data
CITY_STATE_MAPPING = {}
mapping_file = os.path.join(os.path.dirname(__file__), 'static', 'city_to_state_mapping.json')
if os.path.exists(mapping_file):
    with open(mapping_file, 'r', encoding='utf-8') as f:
        raw_mapping = json.load(f)
        # Normalize keys to NFC form for consistent matching
        CITY_STATE_MAPPING = {unicodedata.normalize('NFC', k): v for k, v in raw_mapping.items()}

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False

# Configure caching
app.config['CACHE_TYPE'] = 'SimpleCache'  # In-memory cache
app.config['CACHE_DEFAULT_TIMEOUT'] = 3600  # 1 hour default timeout
cache = Cache(app)

# 数据库配置
DB_CONFIG = {
    'host': 'localhost',
    'database': 'pheno',
    'user': 'postgres',
    'password': '9417941',
    'port': '5432'
}

# pheno_new数据库配置
DB_CONFIG_NEW = {
    'host': 'localhost',
    'database': 'pheno_new',
    'user': 'postgres',
    'password': '9417941',
    'port': '5432'
}

def get_db_connection():
    """获取数据库连接"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"Database connection error: {e}")
        return None

def get_db_connection_new():
    """获取pheno_new数据库连接"""
    try:
        conn = psycopg2.connect(**DB_CONFIG_NEW)
        return conn
    except Exception as e:
        print(f"Database connection error: {e}")
        return None

def dict_fetchall(cursor):
    """将查询结果转换为字典列表"""
    columns = [col[0] for col in cursor.description]
    return [dict(zip(columns, row)) for row in cursor.fetchall()]

def query_by_data_source(data_source, query_pheno, params_pheno, query_new=None, params_new=None):
    """Run a query against pheno, pheno_new, or both databases and return combined results.
    For pheno_new, if query_new is not provided, query_pheno is used.
    """
    results = []
    if data_source in ('pheno', 'both'):
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute(query_pheno, params_pheno)
                results.extend(dict_fetchall(cur))
                cur.close()
                conn.close()
            except Exception as e:
                if conn:
                    conn.close()
                raise e

    if data_source in ('pheno_new', 'both'):
        conn_new = get_db_connection_new()
        if conn_new:
            try:
                cur = conn_new.cursor()
                q = query_new if query_new else query_pheno
                p = params_new if params_new is not None else params_pheno
                cur.execute(q, p)
                results.extend(dict_fetchall(cur))
                cur.close()
                conn_new.close()
            except Exception as e:
                if conn_new:
                    conn_new.close()
                raise e

    return results

@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/new-data')
def new_data():
    """New Data page - opens index with new data modal"""
    return render_template('index.html', open_new_data=True)

@app.route('/geography')
def geography():
    """地理分布页面"""
    return render_template('geography.html')

@app.route('/timeline')
def timeline():
    """时间分析页面"""
    return render_template('timeline.html')

@app.route('/species')
def species_page():
    """物种研究页面"""
    return render_template('species.html')

@app.route('/quality')
def quality():
    """数据质量页面"""
    return render_template('quality.html')

@app.route('/transcription-editor')
def transcription_editor():
    """Transcription file editor page"""
    return render_template('transcription_editor.html')

@app.route('/distribution')
def distribution():
    """数据分布页面"""
    return render_template('distribution.html')

# API 端点
@app.route('/api/overview')
def api_overview():
    """数据概览API"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取基本统计
        stats = {}
        
        # 总观测数
        cursor.execute("SELECT COUNT(*) FROM dwd_observation")
        stats['total_observations'] = cursor.fetchone()[0]
        
        # 站点数
        cursor.execute("SELECT COUNT(*) FROM dwd_station")
        stats['total_stations'] = cursor.fetchone()[0]
        
        # 物种数
        cursor.execute("SELECT COUNT(*) FROM dwd_species")
        stats['total_species'] = cursor.fetchone()[0]
        
        # 物候期数
        cursor.execute("SELECT COUNT(*) FROM dwd_phase")
        stats['total_phases'] = cursor.fetchone()[0]
        
        # 时间范围
        cursor.execute("SELECT MIN(reference_year), MAX(reference_year) FROM dwd_observation")
        year_range = cursor.fetchone()
        pheno_min, pheno_max = year_range[0], year_range[1]

        # 查询 pheno_new 时间范围
        archive_min, archive_max = None, None
        conn_new = get_db_connection_new()
        if conn_new:
            try:
                cur_new = conn_new.cursor()
                cur_new.execute("SELECT MIN(reference_year), MAX(reference_year) FROM dwd_observation")
                new_range = cur_new.fetchone()
                if new_range and new_range[0]:
                    archive_min, archive_max = new_range[0], new_range[1]
                cur_new.close()
                conn_new.close()
            except:
                if conn_new:
                    conn_new.close()

        stats['pheno_basic_range'] = f"{pheno_min}-{pheno_max}"
        if archive_min:
            stats['pheno_archive_range'] = f"{archive_min}-{archive_max}"
        else:
            stats['pheno_archive_range'] = None

        # 最新观测
        cursor.execute("""
            SELECT COUNT(*) FROM dwd_observation
            WHERE reference_year = (SELECT MAX(reference_year) FROM dwd_observation)
        """)
        stats['latest_year_observations'] = cursor.fetchone()[0]

        cursor.close()
        conn.close()
        
        return jsonify(stats)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/stations')
def api_stations():
    """站点数据API - 支持 data_source 参数"""
    data_source = request.args.get('data_source', 'pheno')

    try:
        query_pheno = """
            SELECT
                id, station_name, latitude, longitude,
                altitude, state, area_group, area,
                observation_count
            FROM mv_station_stats
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT
                s.id, s.station_name, s.latitude, s.longitude,
                s.altitude, s.state, s.area_group, s.area,
                COUNT(o.id) as observation_count
            FROM (SELECT DISTINCT ON (id) * FROM dwd_station ORDER BY id) s
            JOIN dwd_observation o ON s.id = o.station_id
            GROUP BY s.id, s.station_name, s.latitude, s.longitude,
                     s.altitude, s.state, s.area_group, s.area
            ORDER BY observation_count DESC
        """
        stations = query_by_data_source(data_source, query_pheno, [], query_new, [])

        # Deduplicate by station_name when combining both sources
        if data_source == 'both':
            seen = {}
            for s in stations:
                name = s['station_name']
                if name not in seen:
                    seen[name] = s
                else:
                    seen[name]['observation_count'] = int(seen[name]['observation_count']) + int(s['observation_count'])
            stations = sorted(seen.values(), key=lambda x: int(x['observation_count']), reverse=True)

        return jsonify(stations)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/species')
def api_species():
    """物种数据API - 支持 data_source 参数"""
    data_source = request.args.get('data_source', 'pheno')

    try:
        query_pheno = """
            SELECT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                sg.group_name,
                s.observation_count
            FROM mv_species_stats s
            LEFT JOIN dwd_species_group sg ON s.id = sg.species_id
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                NULL as group_name,
                COUNT(o.id) as observation_count
            FROM (SELECT DISTINCT ON (id) * FROM dwd_species ORDER BY id) s
            JOIN dwd_observation o ON s.id = o.species_id
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la
            ORDER BY observation_count DESC
        """
        species = query_by_data_source(data_source, query_pheno, [], query_new, [])

        if data_source == 'both':
            seen = {}
            for sp in species:
                sid = sp['id']
                if sid not in seen:
                    seen[sid] = sp
                else:
                    seen[sid]['observation_count'] = int(seen[sid]['observation_count']) + int(sp['observation_count'])
            species = sorted(seen.values(), key=lambda x: int(x['observation_count']), reverse=True)

        return jsonify(species)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/phases')
def api_phases():
    """物候期数据API - 支持 data_source 参数"""
    data_source = request.args.get('data_source', 'pheno')

    try:
        query_pheno = """
            SELECT
                id, phase_name_de, phase_name_en,
                observation_count
            FROM mv_phase_stats
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT
                p.id, p.phase_name_de, p.phase_name_en,
                COUNT(o.id) as observation_count
            FROM (SELECT DISTINCT ON (id) * FROM dwd_phase ORDER BY id) p
            JOIN dwd_observation o ON p.id = o.phase_id
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        phases = query_by_data_source(data_source, query_pheno, [], query_new, [])

        if data_source == 'both':
            seen = {}
            for ph in phases:
                pid = ph['id']
                if pid not in seen:
                    seen[pid] = ph
                else:
                    seen[pid]['observation_count'] = int(seen[pid]['observation_count']) + int(ph['observation_count'])
            phases = sorted(seen.values(), key=lambda x: int(x['observation_count']), reverse=True)

        return jsonify(phases)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/observations')
def api_observations():
    """观测数据API（支持筛选）- supports pheno, pheno_new, or both data sources"""
    data_source = request.args.get('data_source', 'pheno')

    # 获取筛选参数
    station_id = request.args.get('station_id')
    species_id = request.args.get('species_id')
    phase_id = request.args.get('phase_id')
    year_start = request.args.get('year_start')
    year_end = request.args.get('year_end')
    limit = int(request.args.get('limit', 1000))

    def build_obs_query(deduplicate_refs=False):
        where_conditions = []
        params = []

        if station_id:
            where_conditions.append("o.station_id = %s")
            params.append(station_id)
        if species_id:
            where_conditions.append("o.species_id = %s")
            params.append(species_id)
        if phase_id:
            where_conditions.append("o.phase_id = %s")
            params.append(phase_id)
        if year_start:
            where_conditions.append("o.reference_year >= %s")
            params.append(year_start)
        if year_end:
            where_conditions.append("o.reference_year <= %s")
            params.append(year_end)

        where_clause = "WHERE " + " AND ".join(where_conditions) if where_conditions else ""

        # pheno_new has duplicate rows in species/phase/station tables,
        # so use DISTINCT ON subqueries to avoid row multiplication
        if deduplicate_refs:
            query = f"""
                SELECT
                    o.id, o.station_id, o.reference_year, o.species_id,
                    o.phase_id, o.date, o.day_of_year,
                    st.station_name, st.latitude, st.longitude,
                    sp.species_name_en, sp.species_name_de,
                    ph.phase_name_en, ph.phase_name_de
                FROM dwd_observation o
                JOIN (SELECT DISTINCT ON (id) id, station_name, latitude, longitude FROM dwd_station ORDER BY id) st ON o.station_id = st.id
                JOIN (SELECT DISTINCT ON (id) id, species_name_en, species_name_de FROM dwd_species ORDER BY id) sp ON o.species_id = sp.id
                JOIN (SELECT DISTINCT ON (id) id, phase_name_en, phase_name_de FROM dwd_phase ORDER BY id) ph ON o.phase_id = ph.id
                {where_clause}
                ORDER BY o.reference_year DESC, o.day_of_year
                LIMIT %s
            """
        else:
            query = f"""
                SELECT
                    o.id, o.station_id, o.reference_year, o.species_id,
                    o.phase_id, o.date, o.day_of_year,
                    st.station_name, st.latitude, st.longitude,
                    sp.species_name_en, sp.species_name_de,
                    ph.phase_name_en, ph.phase_name_de
                FROM dwd_observation o
                JOIN dwd_station st ON o.station_id = st.id
                JOIN dwd_species sp ON o.species_id = sp.id
                JOIN dwd_phase ph ON o.phase_id = ph.id
                {where_clause}
                ORDER BY o.reference_year DESC, o.day_of_year
                LIMIT %s
            """
        params.append(limit)
        return query, params

    all_observations = []

    if data_source in ('pheno', 'both'):
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            cursor = conn.cursor()
            query, params = build_obs_query()
            cursor.execute(query, params)
            all_observations.extend(dict_fetchall(cursor))
            cursor.close()
            conn.close()
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    if data_source in ('pheno_new', 'both'):
        conn_new = get_db_connection_new()
        if not conn_new:
            return jsonify({'error': 'Pheno_new database connection failed'}), 500
        try:
            cursor_new = conn_new.cursor()
            query, params = build_obs_query(deduplicate_refs=True)
            cursor_new.execute(query, params)
            all_observations.extend(dict_fetchall(cursor_new))
            cursor_new.close()
            conn_new.close()
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # Sort combined results and limit
    all_observations.sort(key=lambda x: (x.get('reference_year', ''), x.get('day_of_year', '')), reverse=True)
    all_observations = all_observations[:limit]

    return jsonify(all_observations)

@app.route('/api/trends')
def api_trends():
    """趋势分析API - supports pheno, pheno_new, or both data sources"""
    data_source = request.args.get('data_source', 'pheno')
    species_id = request.args.get('species_id')
    phase_id = request.args.get('phase_id')
    station_id = request.args.get('station_id')
    year_start = request.args.get('year_start')
    year_end = request.args.get('year_end')

    if not species_id or not phase_id:
        return jsonify({'error': 'species_id and phase_id are required'}), 400

    def build_trends_query(params_list):
        query = """
            SELECT
                reference_year,
                AVG(CAST(day_of_year AS INTEGER)) as avg_day_of_year,
                COUNT(*) as observation_count
            FROM dwd_observation
            WHERE species_id = %s AND phase_id = %s
        """
        params_list.extend([species_id, phase_id])

        if station_id:
            query += " AND station_id = %s"
            params_list.append(station_id)
        if year_start:
            query += " AND reference_year >= %s"
            params_list.append(year_start)
        if year_end:
            query += " AND reference_year <= %s"
            params_list.append(year_end)

        query += """
            GROUP BY reference_year
            ORDER BY reference_year
        """
        return query

    all_trends = []

    if data_source in ('pheno', 'both'):
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            cursor = conn.cursor()
            params = []
            query = build_trends_query(params)
            cursor.execute(query, params)
            all_trends.extend(dict_fetchall(cursor))
            cursor.close()
            conn.close()
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    if data_source in ('pheno_new', 'both'):
        conn_new = get_db_connection_new()
        if not conn_new:
            return jsonify({'error': 'Pheno_new database connection failed'}), 500
        try:
            cursor_new = conn_new.cursor()
            params = []
            query = build_trends_query(params)
            cursor_new.execute(query, params)
            all_trends.extend(dict_fetchall(cursor_new))
            cursor_new.close()
            conn_new.close()
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # If both sources, aggregate by reference_year
    if data_source == 'both' and all_trends:
        aggregated = {}
        for row in all_trends:
            yr = row['reference_year']
            if yr not in aggregated:
                aggregated[yr] = {'reference_year': yr, 'total_day': 0, 'total_count': 0}
            aggregated[yr]['total_day'] += float(row['avg_day_of_year']) * int(row['observation_count'])
            aggregated[yr]['total_count'] += int(row['observation_count'])
        all_trends = [
            {
                'reference_year': v['reference_year'],
                'avg_day_of_year': v['total_day'] / v['total_count'] if v['total_count'] > 0 else 0,
                'observation_count': v['total_count']
            }
            for v in sorted(aggregated.values(), key=lambda x: x['reference_year'])
        ]

    return jsonify(all_trends)

@app.route('/api/quality')
def api_quality():
    """数据质量统计API"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 质量等级分布
        cursor.execute("""
            SELECT 
                ql.id, ql.description,
                COUNT(o.id) as count
            FROM dwd_quality_level ql
            LEFT JOIN dwd_observation o ON ql.id = o.quality_level_id
            GROUP BY ql.id, ql.description
            ORDER BY count DESC
        """)
        
        quality_levels = dict_fetchall(cursor)
        
        # 按年份的质量分布
        cursor.execute("""
            SELECT 
                o.reference_year,
                ql.description,
                COUNT(o.id) as count
            FROM dwd_observation o
            JOIN dwd_quality_level ql ON o.quality_level_id = ql.id
            WHERE o.reference_year >= '1925' AND o.reference_year <= '2020'
            GROUP BY o.reference_year, ql.description
            ORDER BY o.reference_year, ql.description
        """)
        
        quality_by_year = dict_fetchall(cursor)
        
        cursor.close()
        conn.close()
        
        return jsonify({
            'quality_levels': quality_levels,
            'quality_by_year': quality_by_year
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/species-by-phase')
def api_species_by_phase():
    """根据phase搜索species API - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    phase_id = request.args.get('phase_id')
    if not phase_id:
        return jsonify({'error': 'phase_id is required'}), 400

    try:
        query_pheno = """
            SELECT DISTINCT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                sg.group_name,
                COUNT(DISTINCT o.id) as observation_count,
                COUNT(DISTINCT o.station_id) as station_count,
                MIN(o.reference_year) as first_year,
                MAX(o.reference_year) as last_year
            FROM dwd_observation o
            JOIN dwd_species s ON o.species_id = s.id
            LEFT JOIN dwd_species_group sg ON s.id = sg.species_id
            WHERE o.phase_id = %s
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la, sg.group_name
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                NULL as group_name,
                COUNT(DISTINCT o.id) as observation_count,
                COUNT(DISTINCT o.station_id) as station_count,
                MIN(o.reference_year) as first_year,
                MAX(o.reference_year) as last_year
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_species ORDER BY id) s ON o.species_id = s.id
            WHERE o.phase_id = %s
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la
            ORDER BY observation_count DESC
        """
        return jsonify(query_by_data_source(data_source, query_pheno, [phase_id], query_new, [phase_id]))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/species-phases/<species_id>')
def api_species_phases(species_id):
    """获取特定species的所有phases - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')

    try:
        query_pheno = """
            SELECT DISTINCT
                p.id as phase_id,
                p.phase_name_de,
                p.phase_name_en,
                COUNT(o.id) as observation_count,
                MIN(o.reference_year) as first_year,
                MAX(o.reference_year) as last_year,
                AVG(CAST(o.day_of_year AS INTEGER)) as avg_day_of_year
            FROM dwd_observation o
            JOIN dwd_phase p ON o.phase_id = p.id
            WHERE o.species_id = %s
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                p.id as phase_id,
                p.phase_name_de,
                p.phase_name_en,
                COUNT(o.id) as observation_count,
                MIN(o.reference_year) as first_year,
                MAX(o.reference_year) as last_year,
                AVG(CAST(o.day_of_year AS INTEGER)) as avg_day_of_year
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_phase ORDER BY id) p ON o.phase_id = p.id
            WHERE o.species_id = %s
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        return jsonify(query_by_data_source(data_source, query_pheno, [species_id], query_new, [species_id]))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/station-species')
def api_station_species():
    """获取指定站点的所有物种 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    station_id = request.args.get('station_id')
    if not station_id:
        return jsonify({'error': 'station_id is required'}), 400

    try:
        query_pheno = """
            SELECT DISTINCT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_species s ON o.species_id = s.id
            WHERE o.station_id = %s
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_species ORDER BY id) s ON o.species_id = s.id
            WHERE o.station_id = %s
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la
            ORDER BY observation_count DESC
        """
        return jsonify(query_by_data_source(data_source, query_pheno, [station_id], query_new, [station_id]))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/station-phases')
def api_station_phases():
    """获取指定站点的所有物候期 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    station_id = request.args.get('station_id')
    if not station_id:
        return jsonify({'error': 'station_id is required'}), 400

    try:
        query_pheno = """
            SELECT DISTINCT
                p.id, p.phase_name_de, p.phase_name_en,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_phase p ON o.phase_id = p.id
            WHERE o.station_id = %s
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                p.id, p.phase_name_de, p.phase_name_en,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_phase ORDER BY id) p ON o.phase_id = p.id
            WHERE o.station_id = %s
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        return jsonify(query_by_data_source(data_source, query_pheno, [station_id], query_new, [station_id]))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/station-species-phases')
def api_station_species_phases():
    """获取指定站点和物种的所有物候期 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    station_id = request.args.get('station_id')
    species_id = request.args.get('species_id')
    if not station_id or not species_id:
        return jsonify({'error': 'station_id and species_id are required'}), 400

    try:
        query_pheno = """
            SELECT DISTINCT
                p.id as phase_id, p.phase_name_de, p.phase_name_en,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_phase p ON o.phase_id = p.id
            WHERE o.station_id = %s AND o.species_id = %s
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                p.id as phase_id, p.phase_name_de, p.phase_name_en,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_phase ORDER BY id) p ON o.phase_id = p.id
            WHERE o.station_id = %s AND o.species_id = %s
            GROUP BY p.id, p.phase_name_de, p.phase_name_en
            ORDER BY observation_count DESC
        """
        params = [station_id, species_id]
        return jsonify(query_by_data_source(data_source, query_pheno, params, query_new, list(params)))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/station-phase-species')
def api_station_phase_species():
    """获取指定站点和物候期的所有物种 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    station_id = request.args.get('station_id')
    phase_id = request.args.get('phase_id')
    if not station_id or not phase_id:
        return jsonify({'error': 'station_id and phase_id are required'}), 400

    try:
        query_pheno = """
            SELECT DISTINCT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_species s ON o.species_id = s.id
            WHERE o.station_id = %s AND o.phase_id = %s
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                s.id, s.species_name_de, s.species_name_en, s.species_name_la,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_species ORDER BY id) s ON o.species_id = s.id
            WHERE o.station_id = %s AND o.phase_id = %s
            GROUP BY s.id, s.species_name_de, s.species_name_en, s.species_name_la
            ORDER BY observation_count DESC
        """
        params = [station_id, phase_id]
        return jsonify(query_by_data_source(data_source, query_pheno, params, query_new, list(params)))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/species-stations')
def api_species_stations():
    """获取有指定物种观测数据的所有站点 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    species_id = request.args.get('species_id')
    phase_id = request.args.get('phase_id')
    if not species_id:
        return jsonify({'error': 'species_id is required'}), 400

    try:
        where_extra = " AND o.phase_id = %s" if phase_id else ""
        params = [species_id, phase_id] if phase_id else [species_id]

        query_pheno = f"""
            SELECT DISTINCT
                st.id, st.station_name, st.state, st.latitude, st.longitude,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_station st ON o.station_id = st.id
            WHERE o.species_id = %s{where_extra}
            GROUP BY st.id, st.station_name, st.state, st.latitude, st.longitude
            ORDER BY observation_count DESC
        """
        query_new = f"""
            SELECT DISTINCT
                st.id, st.station_name, st.state, st.latitude, st.longitude,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_station ORDER BY id) st ON o.station_id = st.id
            WHERE o.species_id = %s{where_extra}
            GROUP BY st.id, st.station_name, st.state, st.latitude, st.longitude
            ORDER BY observation_count DESC
        """
        return jsonify(query_by_data_source(data_source, query_pheno, params, query_new, list(params)))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/phase-stations')
def api_phase_stations():
    """获取有指定物候期观测数据的所有站点 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    phase_id = request.args.get('phase_id')
    species_id = request.args.get('species_id')
    if not phase_id:
        return jsonify({'error': 'phase_id is required'}), 400

    try:
        where_extra = " AND o.species_id = %s" if species_id else ""
        params = [phase_id, species_id] if species_id else [phase_id]

        query_pheno = f"""
            SELECT DISTINCT
                st.id, st.station_name, st.state, st.latitude, st.longitude,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_station st ON o.station_id = st.id
            WHERE o.phase_id = %s{where_extra}
            GROUP BY st.id, st.station_name, st.state, st.latitude, st.longitude
            ORDER BY observation_count DESC
        """
        query_new = f"""
            SELECT DISTINCT
                st.id, st.station_name, st.state, st.latitude, st.longitude,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_station ORDER BY id) st ON o.station_id = st.id
            WHERE o.phase_id = %s{where_extra}
            GROUP BY st.id, st.station_name, st.state, st.latitude, st.longitude
            ORDER BY observation_count DESC
        """
        return jsonify(query_by_data_source(data_source, query_pheno, params, query_new, list(params)))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/species-phase-stations')
def api_species_phase_stations():
    """获取有指定物种和物候期观测数据的所有站点 - 支持 data_source"""
    data_source = request.args.get('data_source', 'pheno')
    species_id = request.args.get('species_id')
    phase_id = request.args.get('phase_id')
    if not species_id or not phase_id:
        return jsonify({'error': 'species_id and phase_id are required'}), 400

    try:
        query_pheno = """
            SELECT DISTINCT
                st.id, st.station_name, st.state, st.latitude, st.longitude,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN dwd_station st ON o.station_id = st.id
            WHERE o.species_id = %s AND o.phase_id = %s
            GROUP BY st.id, st.station_name, st.state, st.latitude, st.longitude
            ORDER BY observation_count DESC
        """
        query_new = """
            SELECT DISTINCT
                st.id, st.station_name, st.state, st.latitude, st.longitude,
                COUNT(o.id) as observation_count
            FROM dwd_observation o
            JOIN (SELECT DISTINCT ON (id) * FROM dwd_station ORDER BY id) st ON o.station_id = st.id
            WHERE o.species_id = %s AND o.phase_id = %s
            GROUP BY st.id, st.station_name, st.state, st.latitude, st.longitude
            ORDER BY observation_count DESC
        """
        params = [species_id, phase_id]
        return jsonify(query_by_data_source(data_source, query_pheno, params, query_new, list(params)))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pheno-new/species')
def api_pheno_new_species():
    """获取pheno_new数据库中的物种数据，并标记在pheno数据库中是否存在"""
    conn_new = get_db_connection_new()
    conn = get_db_connection()
    if not conn_new or not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    
    try:
        cursor_new = conn_new.cursor()
        cursor = conn.cursor()
        
        # 获取pheno_new中的所有物种及其观测地点
        cursor_new.execute("""
            SELECT 
                s.id as species_id,
                s.species_name_en,
                s.species_name_la,
                s.species_name_de,
                COUNT(DISTINCT o.id) as observation_count,
                STRING_AGG(DISTINCT st.station_name, ', ' ORDER BY st.station_name) as locations
            FROM dwd_species s
            LEFT JOIN dwd_observation o ON s.id = o.species_id
            LEFT JOIN dwd_station st ON o.station_id = st.id
            GROUP BY s.id, s.species_name_en, s.species_name_la, s.species_name_de
            ORDER BY COUNT(DISTINCT o.id) DESC
        """)
        
        new_species = dict_fetchall(cursor_new)
        
        # 获取pheno数据库中的所有物种名称
        cursor.execute("""
            SELECT DISTINCT species_name_de FROM dwd_species
            UNION
            SELECT DISTINCT species_name_en FROM dwd_species
            UNION 
            SELECT DISTINCT species_name_la FROM dwd_species
        """)
        
        existing_species_names = set(row[0] for row in cursor.fetchall() if row[0])
        
        # 标记每个物种是否在pheno数据库中存在
        for species in new_species:
            # 检查任何一个名称是否在pheno数据库中存在
            species['exists_in_pheno'] = (
                species['species_name_en'] in existing_species_names or
                species['species_name_la'] in existing_species_names or
                species['species_name_de'] in existing_species_names
            )
            # 选择一个非空的名称作为显示名称
            species['species_name'] = (
                species['species_name_en'] or 
                species['species_name_la'] or 
                species['species_name_de'] or 
                'Unknown'
            )
            # 标准化locations显示 - 将N/A改为Unknown
            if not species['locations'] or species['locations'] == 'N/A':
                species['locations'] = 'Unknown'
        
        cursor_new.close()
        cursor.close()
        conn_new.close()
        conn.close()
        
        return jsonify(new_species)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pheno-new/species-phases/<species_name>')
def api_pheno_new_species_phases(species_name):
    """获取pheno_new中特定物种的物候期数据，并与pheno数据库中的数据对比"""
    conn_new = get_db_connection_new()
    conn = get_db_connection()
    if not conn_new or not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    
    try:
        cursor_new = conn_new.cursor()
        cursor = conn.cursor()
        
        # 获取pheno_new中该物种的物候期数据
        # 使用子查询避免species表重复记录导致的重复计数
        cursor_new.execute("""
            SELECT DISTINCT
                p.id as phase_id,
                p.phase_name_en,
                p.phase_name_de,
                COUNT(o.id) as observation_count,
                MIN(o.date) as start_date,
                MAX(o.date) as end_date
            FROM dwd_observation o
            JOIN dwd_phase p ON o.phase_id = p.id
            WHERE o.species_id IN (
                SELECT DISTINCT id FROM dwd_species
                WHERE species_name_en = %s
                   OR species_name_la = %s
                   OR species_name_de = %s
            )
            GROUP BY p.id, p.phase_name_en, p.phase_name_de
            ORDER BY p.phase_name_en
        """, (species_name, species_name, species_name))
        
        new_phases = dict_fetchall(cursor_new)
        
        # 获取pheno_new中的时间序列数据
        # 先获取年份范围来决定粒度
        species_filter_sql = """
            SELECT DISTINCT id FROM dwd_species
            WHERE species_name_en = %s OR species_name_la = %s OR species_name_de = %s
        """
        cursor_new.execute(f"""
            SELECT MIN(CAST(o.reference_year AS INTEGER)), MAX(CAST(o.reference_year AS INTEGER))
            FROM dwd_observation o
            WHERE o.species_id IN ({species_filter_sql}) AND o.date IS NOT NULL
        """, (species_name, species_name, species_name))
        year_range = cursor_new.fetchone()
        new_year_span = (year_range[1] - year_range[0] + 1) if year_range and year_range[0] else 0

        if new_year_span <= 10:
            # 10年以内：返回按天去重的观测点
            cursor_new.execute(f"""
                SELECT
                    p.phase_name_en,
                    p.phase_name_de,
                    CAST(o.date AS date) as obs_date,
                    CAST(o.day_of_year AS INTEGER) as day_of_year,
                    CAST(o.reference_year AS INTEGER) as reference_year
                FROM dwd_observation o
                JOIN dwd_phase p ON o.phase_id = p.id
                WHERE o.species_id IN ({species_filter_sql})
                    AND o.date IS NOT NULL
                GROUP BY p.phase_name_en, p.phase_name_de, CAST(o.date AS date), o.day_of_year, o.reference_year
                ORDER BY p.phase_name_en, obs_date
            """, (species_name, species_name, species_name))
            new_time_series = []
            new_individual_observations = dict_fetchall(cursor_new)
        else:
            # 10年以上：按年份聚合
            cursor_new.execute(f"""
                SELECT
                    p.phase_name_en,
                    p.phase_name_de,
                    CAST(o.reference_year AS INTEGER) as year,
                    AVG(CAST(o.day_of_year AS INTEGER)) as avg_day_of_year,
                    COUNT(o.id) as observation_count
                FROM dwd_observation o
                JOIN dwd_phase p ON o.phase_id = p.id
                WHERE o.species_id IN ({species_filter_sql})
                    AND o.date IS NOT NULL
                GROUP BY p.phase_name_en, p.phase_name_de, o.reference_year
                ORDER BY p.phase_name_en, year
            """, (species_name, species_name, species_name))
            new_time_series = dict_fetchall(cursor_new)
            new_individual_observations = []
        
        # 查找该物种在pheno数据库中的对应ID（可能有多个匹配）
        cursor.execute("""
            SELECT id, species_name_de, species_name_en, species_name_la 
            FROM dwd_species 
            WHERE species_name_de = %s OR species_name_en = %s OR species_name_la = %s
        """, (species_name, species_name, species_name))
        
        pheno_species_matches = dict_fetchall(cursor)
        
        pheno_phases = []
        pheno_time_series = []
        pheno_individual_observations = []
        if pheno_species_matches:
            # 获取pheno数据库中的物候期数据
            species_ids = [s['id'] for s in pheno_species_matches]
            placeholders = ','.join(['%s'] * len(species_ids))
            cursor.execute(f"""
                SELECT DISTINCT
                    p.id as phase_id,
                    p.phase_name_de,
                    p.phase_name_en,
                    COUNT(o.id) as observation_count,
                    AVG(CAST(o.day_of_year AS INTEGER)) as avg_day_of_year
                FROM dwd_observation o
                JOIN dwd_phase p ON o.phase_id = p.id
                WHERE o.species_id IN ({placeholders})
                GROUP BY p.id, p.phase_name_de, p.phase_name_en
                ORDER BY p.phase_name_de
            """, species_ids)

            pheno_phases = dict_fetchall(cursor)

            # 获取年份范围
            cursor.execute(f"""
                SELECT MIN(CAST(o.reference_year AS INTEGER)), MAX(CAST(o.reference_year AS INTEGER))
                FROM dwd_observation o
                WHERE o.species_id IN ({placeholders}) AND o.day_of_year IS NOT NULL
            """, species_ids)
            pheno_year_range = cursor.fetchone()
            pheno_year_span = (pheno_year_range[1] - pheno_year_range[0] + 1) if pheno_year_range and pheno_year_range[0] else 0

            if pheno_year_span <= 10:
                # 10年以内：返回按天去重的观测点
                cursor.execute(f"""
                    SELECT
                        p.phase_name_en,
                        p.phase_name_de,
                        CAST(o.date AS date) as obs_date,
                        CAST(o.day_of_year AS INTEGER) as day_of_year,
                        CAST(o.reference_year AS INTEGER) as reference_year
                    FROM dwd_observation o
                    JOIN dwd_phase p ON o.phase_id = p.id
                    WHERE o.species_id IN ({placeholders})
                        AND o.date IS NOT NULL
                    GROUP BY p.phase_name_en, p.phase_name_de, CAST(o.date AS date), o.day_of_year, o.reference_year
                    ORDER BY p.phase_name_en, obs_date
                """, species_ids)
                pheno_individual_observations = dict_fetchall(cursor)
            else:
                # 10年以上：按年份聚合
                cursor.execute(f"""
                    SELECT
                        p.phase_name_en,
                        p.phase_name_de,
                        o.reference_year as year,
                        AVG(CAST(o.day_of_year AS INTEGER)) as avg_day_of_year,
                        COUNT(o.id) as observation_count
                    FROM dwd_observation o
                    JOIN dwd_phase p ON o.phase_id = p.id
                    WHERE o.species_id IN ({placeholders})
                        AND o.day_of_year IS NOT NULL
                    GROUP BY p.phase_name_en, p.phase_name_de, o.reference_year
                    ORDER BY p.phase_name_en, o.reference_year
                """, species_ids)
                pheno_time_series = dict_fetchall(cursor)

        cursor_new.close()
        cursor.close()
        conn_new.close()
        conn.close()

        return jsonify({
            'pheno_new_phases': new_phases,
            'pheno_new_time_series': new_time_series,
            'pheno_new_individual': new_individual_observations,
            'pheno_phases': pheno_phases,
            'pheno_time_series': pheno_time_series,
            'pheno_individual': pheno_individual_observations,
            'pheno_species_matches': pheno_species_matches
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pheno-new/locations')
def api_pheno_new_locations():
    """获取pheno_new数据库中的地理位置并转换为坐标"""
    conn_new = get_db_connection_new()
    if not conn_new:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        cursor_new = conn_new.cursor()

        # Get unique locations and their observation counts from pheno_new
        cursor_new.execute("""
            SELECT
                st.station_name as location,
                COUNT(DISTINCT o.id) as observation_count
            FROM dwd_station st
            INNER JOIN dwd_observation o ON st.id = o.station_id
            WHERE st.area_group = 'Historical'
            GROUP BY st.station_name
            ORDER BY st.station_name
        """)

        locations_data = dict_fetchall(cursor_new)

        # Geocode locations and filter out those without coordinates
        geocoded_locations = []
        for loc in locations_data:
            location_name = loc['location']
            if location_name and not location_name.startswith('Historical Station'):
                geocoded = geocode_location(location_name)
                if geocoded:
                    geocoded['observations'] = loc['observation_count']
                    geocoded_locations.append(geocoded)

        cursor_new.close()
        conn_new.close()

        return jsonify(geocoded_locations)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/data-distribution')
@cache.cached(timeout=7200)  # Cache for 2 hours
def api_data_distribution():
    """获取数据时空分布统计 - 同时从pheno和pheno_new数据库"""
    conn = get_db_connection()
    conn_new = get_db_connection_new()

    if not conn:
        return jsonify({'error': 'Pheno database connection failed'}), 500

    try:
        cursor = conn.cursor()

        # ===== PHENO数据库数据 (using materialized views for speed) =====
        # 获取年份-地区的观测数量分布（按1年）
        cursor.execute("""
            SELECT year, state, observation_count
            FROM mv_year_state_distribution
            ORDER BY year, state
        """)

        pheno_time_location_dist = dict_fetchall(cursor)

        # 获取月份分布
        cursor.execute("""
            SELECT year, month, observation_count
            FROM mv_year_month_distribution
            ORDER BY year, month
        """)

        pheno_month_dist = dict_fetchall(cursor)

        # 获取数据覆盖范围统计
        cursor.execute("""
            SELECT min_year, max_year, station_count, species_count, phase_count
            FROM mv_coverage_stats
        """)

        pheno_coverage = dict_fetchone(cursor)

        cursor.close()
        conn.close()

        # ===== PHENO_NEW数据库数据 =====
        pheno_new_time_location_dist = []
        pheno_new_month_dist = []
        pheno_new_coverage = None

        if conn_new:
            cursor_new = conn_new.cursor()

            # 获取年份-地区的观测数量分布（按1年）
            # For pheno_new, use station_name as location and map to state
            # Use DISTINCT to avoid duplicate counting due to duplicate station records
            cursor_new.execute("""
                SELECT
                    CAST(reference_year AS INTEGER) as year,
                    s.station_name,
                    COUNT(DISTINCT o.id) as observation_count
                FROM dwd_observation o
                JOIN dwd_station s ON o.station_id = s.id
                WHERE s.station_name IS NOT NULL
                  AND s.station_name != ''
                  AND s.area_group = 'Historical'
                  AND NOT s.station_name LIKE 'Historical Station%'
                GROUP BY CAST(reference_year AS INTEGER), s.station_name
                ORDER BY year, s.station_name
            """)

            raw_pheno_new_data = dict_fetchall(cursor_new)

            # Map city names to state names and aggregate
            state_aggregated = {}
            for item in raw_pheno_new_data:
                city_name = item['station_name']
                # Normalize city name to NFC form for consistent matching
                city_name_normalized = unicodedata.normalize('NFC', city_name)
                state_name = CITY_STATE_MAPPING.get(city_name_normalized, city_name)  # Use mapping or keep original
                year = item['year']
                count = item['observation_count']

                key = (year, state_name)
                if key not in state_aggregated:
                    state_aggregated[key] = 0
                state_aggregated[key] += count

            # Convert back to list format
            pheno_new_time_location_dist = [
                {
                    'year': year,
                    'state': state,
                    'observation_count': count
                }
                for (year, state), count in sorted(state_aggregated.items())
            ]

            # 获取月份分布
            cursor_new.execute("""
                SELECT
                    CAST(reference_year AS INTEGER) as year,
                    CASE
                        WHEN CAST(day_of_year AS INTEGER) <= 31 THEN 1
                        WHEN CAST(day_of_year AS INTEGER) <= 59 THEN 2
                        WHEN CAST(day_of_year AS INTEGER) <= 90 THEN 3
                        WHEN CAST(day_of_year AS INTEGER) <= 120 THEN 4
                        WHEN CAST(day_of_year AS INTEGER) <= 151 THEN 5
                        WHEN CAST(day_of_year AS INTEGER) <= 181 THEN 6
                        WHEN CAST(day_of_year AS INTEGER) <= 212 THEN 7
                        WHEN CAST(day_of_year AS INTEGER) <= 243 THEN 8
                        WHEN CAST(day_of_year AS INTEGER) <= 273 THEN 9
                        WHEN CAST(day_of_year AS INTEGER) <= 304 THEN 10
                        WHEN CAST(day_of_year AS INTEGER) <= 334 THEN 11
                        ELSE 12
                    END as month,
                    COUNT(o.id) as observation_count
                FROM dwd_observation o
                WHERE day_of_year IS NOT NULL AND day_of_year != ''
                GROUP BY CAST(reference_year AS INTEGER), month
                ORDER BY year, month
            """)

            pheno_new_month_dist = dict_fetchall(cursor_new)

            # 获取数据覆盖范围统计 (count unique station names, not IDs,
            # because old import created multiple IDs per physical station)
            cursor_new.execute("""
                SELECT
                    MIN(CAST(o.reference_year AS INTEGER)) as min_year,
                    MAX(CAST(o.reference_year AS INTEGER)) as max_year,
                    COUNT(DISTINCT s.station_name) as station_count,
                    COUNT(DISTINCT o.species_id) as species_count,
                    COUNT(DISTINCT o.phase_id) as phase_count
                FROM dwd_observation o
                JOIN dwd_station s ON o.station_id = s.id
            """)

            pheno_new_coverage = dict_fetchone(cursor_new)

            cursor_new.close()
            conn_new.close()

        return jsonify({
            'pheno': {
                'time_location_distribution': pheno_time_location_dist,
                'month_distribution': pheno_month_dist,
                'coverage': pheno_coverage
            },
            'pheno_new': {
                'time_location_distribution': pheno_new_time_location_dist,
                'month_distribution': pheno_new_month_dist,
                'coverage': pheno_new_coverage
            }
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

def dict_fetchone(cursor):
    """将单行查询结果转换为字典"""
    columns = [col[0] for col in cursor.description]
    row = cursor.fetchone()
    return dict(zip(columns, row)) if row else None

@app.route('/api/debug/pheno-new-stations')
def api_debug_pheno_new_stations():
    """Debug endpoint to check pheno_new station data"""
    conn_new = get_db_connection_new()

    if not conn_new:
        return jsonify({'error': 'Pheno_new database connection failed'}), 500

    try:
        cursor_new = conn_new.cursor()

        # Check station table structure
        cursor_new.execute("""
            SELECT column_name, data_type
            FROM information_schema.columns
            WHERE table_name = 'dwd_station'
            ORDER BY ordinal_position
        """)

        columns = dict_fetchall(cursor_new)

        # Get sample station data
        cursor_new.execute("""
            SELECT * FROM dwd_station LIMIT 5
        """)

        sample_stations = dict_fetchall(cursor_new)

        # Get stations with coordinates
        cursor_new.execute("""
            SELECT COUNT(*) as total_stations,
                   COUNT(latitude) as stations_with_lat,
                   COUNT(longitude) as stations_with_lon
            FROM dwd_station
        """)

        coord_stats = dict_fetchone(cursor_new)

        cursor_new.close()
        conn_new.close()

        return jsonify({
            'columns': columns,
            'sample_stations': sample_stations,
            'coordinate_stats': coord_stats
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/data-distribution-detailed')
@cache.cached(timeout=7200)  # Cache for 2 hours
def api_data_distribution_detailed():
    """获取详细的站点级别数据分布 - 用于地图和时间线可视化"""
    conn = get_db_connection()
    conn_new = get_db_connection_new()

    if not conn:
        return jsonify({'error': 'Pheno database connection failed'}), 500

    try:
        cursor = conn.cursor()

        # ===== PHENO数据库 - 站点级别聚合（使用物化视图优化） =====
        cursor.execute("""
            SELECT
                station_id,
                station_name,
                latitude,
                longitude,
                state,
                area,
                reference_year,
                observation_count
            FROM mv_station_yearly_stats
            ORDER BY station_name, reference_year
        """)

        pheno_station_yearly = dict_fetchall(cursor)

        cursor.close()
        conn.close()

        # ===== PHENO_NEW数据库 - 站点级别聚合 =====
        pheno_new_station_yearly = []

        if conn_new:
            cursor_new = conn_new.cursor()

            # Coordinates are now stored in DB for all stations
            cursor_new.execute("""
                SELECT
                    s.id as station_id,
                    s.station_name,
                    s.latitude,
                    s.longitude,
                    s.state,
                    s.area,
                    o.reference_year,
                    COUNT(o.id) as observation_count
                FROM dwd_observation o
                JOIN dwd_station s ON o.station_id = s.id
                WHERE s.station_name IS NOT NULL
                  AND s.latitude IS NOT NULL
                  AND s.longitude IS NOT NULL
                GROUP BY s.id, s.station_name, s.latitude, s.longitude, s.state, s.area, o.reference_year
                ORDER BY s.station_name, o.reference_year
            """)

            pheno_new_station_yearly = dict_fetchall(cursor_new)

            cursor_new.close()
            conn_new.close()

        return jsonify({
            'pheno': pheno_station_yearly,
            'pheno_new': pheno_new_station_yearly
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Transcription Editor API endpoints
TRANSCRIPTION_BASE_PATH = os.environ.get('TRANSCRIPTION_PATH', os.path.join(os.path.dirname(os.path.abspath(__file__)), 'corrected'))

@app.route('/api/transcription/folders')
def api_transcription_folders():
    """Get all folders list"""
    try:
        folders = []
        base_path = Path(TRANSCRIPTION_BASE_PATH)
        
        for folder in sorted(base_path.iterdir()):
            if folder.is_dir():
                # Extract folder index
                match = re.match(r'^(\d+)', folder.name)
                index = match.group(1) if match else None
                
                # Check for files (odt can be direct .odt or .zip containing .odt)
                files_in_folder = [f for f in folder.iterdir() if f.is_file()]
                has_odt = any(f.suffix == '.odt' for f in files_in_folder) or \
                          any(f.suffix == '.zip' for f in files_in_folder)
                has_tif = any(f.suffix == '.tif' for f in files_in_folder)
                is_tabelle = 'Tabelle' in folder.name
                
                folders.append({
                    'name': folder.name,
                    'index': index,
                    'hasOdt': has_odt,
                    'hasTif': has_tif,
                    'isTabelle': is_tabelle
                })
        
        return jsonify(folders)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/transcription/folder/<path:folder_name>')
def api_transcription_folder_contents(folder_name):
    """Get folder contents"""
    try:
        folder_path = Path(TRANSCRIPTION_BASE_PATH) / folder_name
        if not folder_path.exists():
            return jsonify({'error': 'Folder not found'}), 404
        
        files = []
        for file in sorted(folder_path.iterdir()):
            if file.is_file():
                if file.suffix in ['.odt', '.tif']:
                    files.append({
                        'name': file.name,
                        'type': file.suffix[1:]  # Remove the dot
                    })
                elif file.suffix == '.zip':
                    # Treat zip files containing odt as odt files
                    files.append({
                        'name': file.name,
                        'type': 'odt'
                    })
        
        return jsonify({'files': files})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/transcription/image/<path:folder_name>/<path:file_name>')
def api_transcription_image(folder_name, file_name):
    """Get TIF image"""
    try:
        file_path = Path(TRANSCRIPTION_BASE_PATH) / folder_name / file_name
        if not file_path.exists():
            return jsonify({'error': 'File not found'}), 404
        
        # Convert TIF to PNG for web display
        img = Image.open(file_path)
        img_io = io.BytesIO()
        img.save(img_io, 'PNG')
        img_io.seek(0)
        
        return send_file(img_io, mimetype='image/png')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def _extract_odt_tables(file_path):
    """Extract tables and text from an ODT file"""
    doc = load(str(file_path))
    tables = []
    raw_content = []

    for p in doc.getElementsByType(text.P):
        raw_content.append(teletype.extractText(p))

    for table in doc.getElementsByType(Table):
        table_data = []
        for row in table.getElementsByType(TableRow):
            row_data = []
            for cell in row.getElementsByType(TableCell):
                cell_text = ""
                for p in cell.getElementsByType(text.P):
                    cell_text += teletype.extractText(p)

                repeat = cell.getAttribute("numbercolumnsrepeated")
                repeat_count = int(repeat) if repeat else 1
                colspan = cell.getAttribute("numbercolumnsspanned")
                colspan_count = int(colspan) if colspan else 1
                rowspan = cell.getAttribute("numberrowsspanned")
                rowspan_count = int(rowspan) if rowspan else 1

                for _ in range(repeat_count):
                    if colspan_count > 1 or rowspan_count > 1:
                        cell_obj = {'text': cell_text.strip()}
                        if colspan_count > 1:
                            cell_obj['colspan'] = colspan_count
                        if rowspan_count > 1:
                            cell_obj['rowspan'] = rowspan_count
                        row_data.append(cell_obj)
                    else:
                        row_data.append(cell_text.strip())

            if row_data:
                table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    return tables, '\n'.join(raw_content)


def _extract_docx_tables(file_path):
    """Extract tables and text from a DOCX file"""
    doc = DocxDocument(str(file_path))
    tables = []
    raw_content = []

    for para in doc.paragraphs:
        if para.text.strip():
            raw_content.append(para.text.strip())

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            seen_tcs = set()
            for cell in row.cells:
                tc = cell._tc
                tc_id = id(tc)
                if tc_id in seen_tcs:
                    continue
                seen_tcs.add(tc_id)

                cell_text = cell.text.strip()
                tc_pr = tc.find(qn('w:tcPr'))
                colspan = 1
                rowspan = 1
                if tc_pr is not None:
                    gs = tc_pr.find(qn('w:gridSpan'))
                    if gs is not None:
                        colspan = int(gs.get(qn('w:val'), '1'))
                    vm = tc_pr.find(qn('w:vMerge'))
                    if vm is not None:
                        val = vm.get(qn('w:val'), '')
                        if val == 'restart':
                            rowspan = 2
                        else:
                            continue  # Skip continuation cells

                if colspan > 1 or rowspan > 1:
                    cell_obj = {'text': cell_text}
                    if colspan > 1:
                        cell_obj['colspan'] = colspan
                    if rowspan > 1:
                        cell_obj['rowspan'] = rowspan
                    row_data.append(cell_obj)
                else:
                    row_data.append(cell_text)

            if row_data:
                table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    return tables, '\n'.join(raw_content)


@app.route('/api/transcription/odt/<path:folder_name>/<path:file_name>')
def api_transcription_odt(folder_name, file_name):
    """Get document content (supports .odt, .docx, and .zip containing either)"""
    try:
        file_path = Path(TRANSCRIPTION_BASE_PATH) / folder_name / file_name
        if not file_path.exists():
            return jsonify({'error': 'File not found'}), 404

        if file_path.suffix == '.zip':
            with zipfile.ZipFile(str(file_path), 'r') as zf:
                odt_names = [n for n in zf.namelist() if n.endswith('.odt')]
                docx_names = [n for n in zf.namelist() if n.endswith('.docx')]
                if odt_names:
                    with tempfile.NamedTemporaryFile(suffix='.odt', delete=False) as tmp:
                        tmp.write(zf.read(odt_names[0]))
                        tmp_path = tmp.name
                    tables, raw_content = _extract_odt_tables(tmp_path)
                    os.unlink(tmp_path)
                elif docx_names:
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        tmp.write(zf.read(docx_names[0]))
                        tmp_path = tmp.name
                    tables, raw_content = _extract_docx_tables(tmp_path)
                    os.unlink(tmp_path)
                else:
                    return jsonify({'error': 'No ODT or DOCX file found in zip'}), 404
        elif file_path.suffix == '.docx':
            tables, raw_content = _extract_docx_tables(file_path)
        else:
            tables, raw_content = _extract_odt_tables(file_path)

        return jsonify({
            'raw_content': raw_content,
            'tables': tables
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/transcription/annotations/<path:folder_name>/<path:file_name>')
def api_transcription_annotations(folder_name, file_name):
    """Get annotations for a specific file"""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500

        cursor = conn.cursor()

        # Create table if not exists
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS transcription_annotations (
                id SERIAL PRIMARY KEY,
                folder_name VARCHAR(500) NOT NULL,
                file_name VARCHAR(500) NOT NULL,
                annotation_text TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.commit()

        # Get annotations for this file
        cursor.execute("""
            SELECT id, folder_name, file_name, annotation_text, created_at
            FROM transcription_annotations
            WHERE folder_name = %s AND file_name = %s
            ORDER BY created_at DESC
        """, (folder_name, file_name))

        annotations = dict_fetchall(cursor)

        cursor.close()
        conn.close()

        return jsonify({'annotations': annotations})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/transcription/annotations', methods=['POST'])
def api_transcription_add_annotation():
    """Add a new annotation"""
    try:
        data = request.json
        folder_name = data.get('folder_name')
        file_name = data.get('file_name')
        annotation_text = data.get('annotation_text')

        if not folder_name or not file_name or not annotation_text:
            return jsonify({'error': 'Missing required fields'}), 400

        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500

        cursor = conn.cursor()

        # Create table if not exists
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS transcription_annotations (
                id SERIAL PRIMARY KEY,
                folder_name VARCHAR(500) NOT NULL,
                file_name VARCHAR(500) NOT NULL,
                annotation_text TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Insert annotation
        cursor.execute("""
            INSERT INTO transcription_annotations (folder_name, file_name, annotation_text)
            VALUES (%s, %s, %s)
            RETURNING id, created_at
        """, (folder_name, file_name, annotation_text))

        result = cursor.fetchone()
        conn.commit()

        cursor.close()
        conn.close()

        return jsonify({
            'success': True,
            'message': 'Annotation added successfully',
            'id': result[0],
            'created_at': result[1].isoformat()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/species-mapping')
def api_species_mapping():
    """Read species mapping from final_species_mapping.csv and return as JSON"""
    try:
        csv_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'final_species_mapping.csv')
        if not os.path.exists(csv_path):
            return jsonify({'error': 'Species mapping file not found'}), 404

        rows = []
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(row)

        return jsonify({'data': rows})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0',port=9090)